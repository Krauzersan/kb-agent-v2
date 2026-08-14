"""Извлечение текста из файлов разных типов, нарезка на куски и индексация."""
from __future__ import annotations

import csv
import hashlib
import io
import ipaddress
import logging
import os
import re
import socket
import sys
from html.parser import HTMLParser
from typing import List
from urllib.parse import urljoin, urlparse

import httpx

import db
import llm
import progress
import settings_store
import vectorstore
from config import settings

log = logging.getLogger("ingest")

TEXT_EXT = {".md", ".markdown", ".txt", ".yml", ".yaml", ".json", ".log", ".csv", ".tsv"}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".gif"}

# Каталог: краткое summary файла для агрегационных вопросов («перечисли все X»).
# Обрезаем вход — нужна не точность до символа, а полнота перечисления сущностей.
_SUMMARY_SYSTEM = (
    "Ты строишь оглавление базы знаний компании. По тексту документа кратко опиши, о чём он, "
    "и ОБЯЗАТЕЛЬНО перечисли поимённо все конкретные названия систем/продуктов/интеграций/касс/"
    "CRM/сервисов, которые в нём упоминаются как предмет статьи (не мимоходом). "
    "Если таких названий нет — просто опиши тему в 1-2 предложениях. "
    "Формат: 1-2 предложения о теме, затем при наличии — строка «Упоминаются: A, B, C». "
    "Не пиши вступлений и не выдумывай названий, которых нет в тексте."
)
_SUMMARY_MAX_INPUT = 6000


def generate_summary(filename: str, text: str) -> str:
    """Короткое summary файла для каталога (используется в rag._answer_from_catalog)."""
    text = (text or "").strip()
    if not text:
        return ""
    user = f"Имя файла: {filename}\n\nТекст документа:\n{text[:_SUMMARY_MAX_INPUT]}"
    try:
        return llm.complete(_SUMMARY_SYSTEM, user, max_tokens=250)
    except llm.LLMNotConfigured:
        return ""
    except Exception:
        log.exception("Не удалось построить summary для файла %s", filename)
        return ""

# Снимаем лимит на размер ячейки CSV (по умолчанию 128 КБ) — для больших полей.
_max_field = sys.maxsize
while True:
    try:
        csv.field_size_limit(_max_field)
        break
    except OverflowError:
        _max_field //= 10


# ---------- извлечение текста ----------

def _read_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_csv(path: str) -> str:
    """CSV/TSV -> читаемые строки 'колонка: значение'."""
    delimiter = "\t" if path.lower().endswith(".tsv") else ","
    out: List[str] = []
    with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f, delimiter=delimiter)
        rows = list(reader)
    if not rows:
        return ""
    header = rows[0]
    for row in rows[1:]:
        pairs = [f"{h}: {v}" for h, v in zip(header, row) if v.strip()]
        if pairs:
            out.append("; ".join(pairs))
    # если структура нестандартная — отдаём как есть
    return "\n".join(out) if out else "\n".join(",".join(r) for r in rows)


_DOCX_EMBED_RE = re.compile(r'r:embed="(rId\d+)"')
_DOCX_DESCR_RE = re.compile(r'descr="([^"]*)"')


def _save_docx_image(blob: bytes, ext: str) -> str:
    """Кладёт картинку из docx в settings.ASSETS_DIR (та же публичная папка, что и
    для ZIP-импорта) и возвращает публичный URL. Имя — хеш содержимого, поэтому
    повторная переиндексация того же файла не плодит копии."""
    if ext not in IMAGE_EXT:
        ext = ".png"
    dest_rel = f"{hashlib.sha1(blob).hexdigest()[:16]}{ext}"
    dest_abs = os.path.join(settings.ASSETS_DIR, dest_rel)
    os.makedirs(settings.ASSETS_DIR, exist_ok=True)
    if not os.path.exists(dest_abs):
        with open(dest_abs, "wb") as f:
            f.write(blob)
    base_url = (settings_store.get("public_base_url") or "").rstrip("/")
    return f"{base_url}/assets/{dest_rel}"


def _docx_paragraph_images(p, rels) -> List[str]:
    """Встроенные в параграф картинки (docx хранит их как r:embed-ссылки на
    word/media/... внутри XML параграфа) -> маркеры [изображение] URL, в том же
    формате, что и остальной конвейер (ZIP-импорт, парсинг веб-страниц)."""
    xml = p._p.xml
    rids = _DOCX_EMBED_RE.findall(xml)
    if not rids:
        return []
    descrs = _DOCX_DESCR_RE.findall(xml)
    # подпись берём, только если картинка одна и alt-текст задан однозначно —
    # при нескольких картинках в параграфе сопоставить подпись с конкретной
    # ненадёжно, лучше без подписи, чем с неверной
    caption = descrs[0].strip() if len(rids) == 1 and descrs and descrs[0].strip() else ""
    markers = []
    for rid in rids:
        rel = rels.get(rid)
        if rel is None or "image" not in rel.reltype:
            continue
        try:
            ext = os.path.splitext(rel.target_part.partname)[1].lower()
            url = _save_docx_image(rel.target_part.blob, ext)
        except Exception:
            log.exception("Не удалось извлечь картинку из docx (rId=%s)", rid)
            continue
        label = f": {caption}" if caption else ""
        markers.append(f"\n[изображение{label}] {url}\n")
    return markers


def _read_docx(path: str) -> str:
    import docx  # python-docx

    document = docx.Document(path)
    rels = document.part.rels
    parts: List[str] = []
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text)
        parts.extend(_docx_paragraph_images(p, rels))
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_xlsx(path: str) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    out: List[str] = []
    for ws in wb.worksheets:
        out.append(f"# Лист: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def _read_image_ocr(path: str) -> str:
    """OCR картинки (скриншоты, схемы) через Tesseract. Языки: рус + англ."""
    try:
        import pytesseract
        from PIL import Image

        return pytesseract.image_to_string(Image.open(path), lang="rus+eng")
    except Exception as e:  # tesseract не установлен / не распознал
        return f"[Не удалось распознать текст на изображении: {e}]"


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in {".docx"}:
        return _read_docx(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext in {".xlsx", ".xlsm"}:
        return _read_xlsx(path)
    if ext in {".csv", ".tsv"}:
        return _read_csv(path)
    if ext in IMAGE_EXT:
        return _read_image_ocr(path)
    if ext in TEXT_EXT:
        return _read_text(path)
    # неизвестный тип — пробуем как текст
    try:
        return _read_text(path)
    except Exception:
        return ""


# ---------- нарезка на куски ----------

# Заголовок markdown (# ... ###### ...) или нумерованный пункт вида «1.», «2.1.» —
# сигнал структурной границы документа. Требуем ЛИТЕРАЛЬНУЮ точку после номера
# («1. Текст», не «1.2.3 Текст» без точки и не «1 касса подключена» — обычное
# предложение, начинающееся с цифры) — намеренно консервативно: лучше пропустить
# часть настоящих заголовков, чем ложно резать обычный текст по случайным цифрам.
_HEADING_RE = re.compile(r"^(#{1,6}\s+\S.*|\d{1,2}(?:\.\d{1,2}){0,3}\.\s+\S.*)$")


def _split_sections(text: str) -> List[str]:
    """Режет текст на секции по заголовкам/нумерованным пунктам, чтобы упаковка в
    чанки ниже не резала инструкцию (шаги настройки и т.п.) посередине по случайной
    границе строки. Секция — заголовок и всё до следующего такого же сигнала.
    Если структуры нет вообще — вернёт весь текст одной секцией, и дальше отработает
    прежняя построчная упаковка (полная совместимость со старым поведением)."""
    lines = text.split("\n")
    sections: List[str] = []
    buf: List[str] = []
    for line in lines:
        if _HEADING_RE.match(line.strip()) and buf:
            sections.append("\n".join(buf))
            buf = [line]
        else:
            buf.append(line)
    if buf:
        sections.append("\n".join(buf))
    return sections if sections else [text]


def _pack_lines(section: str, size: int, overlap: int, chunks: List[str], buf: str) -> str:
    """Старая построчная упаковка (fallback) — применяется только внутри секции,
    которая сама целиком не влезает в size."""
    paragraphs = [p.strip() for p in section.split("\n") if p.strip()]
    for p in paragraphs:
        if len(buf) + len(p) + 1 <= size:
            buf = f"{buf}\n{p}" if buf else p
        else:
            if buf:
                chunks.append(buf)
            if len(p) <= size:
                buf = p
            else:
                # очень длинный абзац — режем окнами с перекрытием
                start = 0
                while start < len(p):
                    chunks.append(p[start:start + size])
                    start += size - overlap
                buf = ""
    return buf


def chunk_text(text: str) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    size = settings.CHUNK_SIZE
    overlap = settings.CHUNK_OVERLAP
    chunks: List[str] = []
    buf = ""
    for section in _split_sections(text):
        section = section.strip("\n")
        if not section.strip():
            continue
        if len(section) <= size:
            # секция целиком умещается — пакуем как атомарный кусок (не разрываем
            # заголовок/пункт и его содержимое построчно)
            if len(buf) + len(section) + 1 <= size:
                buf = f"{buf}\n{section}" if buf else section
            else:
                if buf:
                    chunks.append(buf)
                buf = section
        else:
            # секция сама больше size — прежняя построчная упаковка внутри неё
            if buf:
                chunks.append(buf)
                buf = ""
            buf = _pack_lines(section, size, overlap, chunks, buf)
    if buf:
        chunks.append(buf)
    return chunks


# ---------- парсинг веб-страниц (по ссылке) ----------

class _HtmlExtract(HTMLParser):
    """Достаёт из HTML читаемый текст и ссылки на изображения (в абсолютном виде)."""

    _SKIP = {"script", "style", "noscript", "svg"}
    _BREAK = {"p", "br", "div", "li", "tr", "h1", "h2", "h3", "h4", "h5", "section"}

    def __init__(self, base_url: str):
        super().__init__()
        self.base = base_url
        self.parts: List[str] = []
        self.title = ""
        self._skip = 0
        self._in_title = False

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip += 1
        elif tag == "title":
            self._in_title = True
        elif tag in self._BREAK:
            self.parts.append("\n")
        elif tag == "img" and not self._skip:
            d = dict(attrs)
            src = d.get("src") or d.get("data-src") or d.get("data-original")
            alt = (d.get("alt") or "").strip()
            if src and not src.startswith("data:"):
                url = urljoin(self.base, src)
                self.parts.append(f"\n[изображение{': ' + alt if alt else ''}] {url}\n")

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip > 0:
            self._skip -= 1
        elif tag == "title":
            self._in_title = False

    def handle_data(self, data):
        if self._skip:
            return
        if self._in_title:
            self.title += data
        else:
            t = data.strip()
            if t:
                self.parts.append(t + " ")


def _assert_public_url(url: str) -> None:
    """Защита от SSRF: не даём сходить на localhost/внутреннюю сеть/link-local
    ни в исходном URL, ни через редирект (проверяется на каждом хопе)."""
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError(f"Недопустимая схема ссылки: {parsed.scheme!r}")
    host = parsed.hostname
    if not host:
        raise ValueError("Не удалось определить хост ссылки")
    try:
        infos = socket.getaddrinfo(host, None)
    except socket.gaierror as e:
        raise ValueError(f"Не удалось разрешить хост {host!r}: {e}") from e
    for info in infos:
        ip = ipaddress.ip_address(info[4][0])
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved or ip.is_multicast:
            raise ValueError(f"Ссылка ведёт на внутренний/служебный адрес ({ip}) — запрещено")


def fetch_url_text(url: str):
    """Скачивает страницу и возвращает (заголовок, текст с ссылками на картинки)."""
    hops = 0
    next_url = url
    with httpx.Client(follow_redirects=False, timeout=30,
                       headers={"User-Agent": "Mozilla/5.0 (kb-agent)"}) as client:
        while True:
            _assert_public_url(next_url)
            r = client.get(next_url)
            if r.is_redirect and hops < 5:
                next_url = str(r.next_request.url)
                hops += 1
                continue
            break
    r.raise_for_status()
    p = _HtmlExtract(str(r.url))
    p.feed(r.text)
    text = re.sub(r"[ \t]+\n", "\n", "".join(p.parts))
    text = re.sub(r"\n\s*\n+", "\n\n", text).strip()
    title = re.sub(r"\s+", " ", (p.title or url)).strip()[:120] or url
    return title, text


# ---------- основной конвейер ----------

def ingest_file(file_id: int, skip_summary: bool = False) -> None:
    """Индексирует один файл по его id. Вызывается в фоне после загрузки.

    skip_summary=True — не пересчитывать summary для каталога (без вызовов LLM).
    Нужно для лёгкой массовой переиндексации ТОЛЬКО эмбеддингов (например после
    изменения того, что в них попадает) — иначе на 1000+ файлов пришлось бы зря
    пересчитывать уже готовые summary через API.
    """
    rec = db.get_file(file_id)
    if not rec:
        return
    db.set_status(file_id, "indexing")
    try:
        vectorstore.ensure_collection()
        text = extract_text(rec["stored_path"])
        chunks = chunk_text(text)
        # прогресс по кускам текущего файла (двигает бар даже на тяжёлом файле)
        progress.set_file(rec["filename"], len(chunks))
        # на случай переиндексации — убираем старые куски этого файла
        vectorstore.delete_file(file_id)
        db.delete_chunks_fts(file_id)
        n = vectorstore.add_chunks(file_id, rec["filename"], chunks,
                                   on_progress=progress.chunk_step,
                                   priority=rec.get("priority", 0))
        if n == 0:
            db.set_status(file_id, "error", chunks=0,
                          error="Не удалось извлечь текст (пустой файл или нераспознан).")
        else:
            db.index_chunks_fts(file_id, chunks)   # лексический индекс — для гибридного поиска
            db.set_status(file_id, "ready", chunks=n, error=None)
            if skip_summary:
                return
            # Каталог для агрегационных вопросов — best-effort, не должен ронять индексацию.
            try:
                summary = generate_summary(rec["filename"], text)
                if summary:
                    db.set_summary(file_id, summary)
            except Exception:
                log.exception("Индексация: не удалось построить summary для %s", rec["filename"])
    except Exception as e:  # noqa: BLE001
        db.set_status(file_id, "error", error=str(e))
